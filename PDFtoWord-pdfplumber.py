# pip install pdfplumber python-docx
# python PDFtoWord.py

# Comment: this library cannot export table in PDF to word, worse than PyMuPDF and pdf2docx
# pdf2docx > PyMuPDF > pdfplumber

import pdfplumber
from docx import Document

def extract_text_and_tables(pdf_path):
    text = ""
    tables = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text += page.extract_text() or ""

            # Extract tables
            page_tables = page.extract_tables()
            for table in page_tables:
                tables.append(table)

    return text, tables

def create_docx(text, tables, output_docx_path):
    doc = Document()
    doc.add_paragraph(text)

    for table in tables:
        if table:
            # Create a table with the correct number of columns
            num_cols = len(table[0])
            doc_table = doc.add_table(rows=0, cols=num_cols)
            for row_data in table:
                row_cells = doc_table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = str(cell_data)

    doc.save(output_docx_path)

def convert_pdf_to_docx(pdf_path, output_docx_path):
    text, tables = extract_text_and_tables(pdf_path)
    create_docx(text, tables, output_docx_path)

if __name__ == "__main__":
    input_pdf_path = 'input.pdf'
    output_docx_path = 'output.docx'
    convert_pdf_to_docx(input_pdf_path, output_docx_path)
