# pip install PyMuPDF python-docx
# python PDFtoWord.py

# Comment: this library cannot export table in PDF to word, worse than pdf2docx, better than pdfplumber
# pdf2docx > PyMuPDF > pdfplumber

import fitz  # PyMuPDF
from docx import Document

def extract_text_and_tables(pdf_path):
    doc = fitz.open(pdf_path)
    text = ""
    tables = []

    for page in doc:
        text += page.get_text()

        # Extract tables (this is a simplified example)
        # You might need to use a more sophisticated method to extract tables
        for block in page.get_text("dict")["blocks"]:
            if "lines" in block:
                table = []
                for line in block["lines"]:
                    row = [span["text"] for span in line["spans"]]
                    table.append(row)
                tables.append(table)

    return text, tables

def create_docx(text, tables, output_docx_path):
    doc = Document()
    doc.add_paragraph(text)

    for table in tables:
        if table:
            # Create a table with the correct number of columns
            num_cols = max(len(row) for row in table)
            doc_table = doc.add_table(rows=0, cols=num_cols)
            for row_data in table:
                row_cells = doc_table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    if i < len(row_cells):
                        row_cells[i].text = cell_data

    doc.save(output_docx_path)

def convert_pdf_to_docx(pdf_path, output_docx_path):
    text, tables = extract_text_and_tables(pdf_path)
    create_docx(text, tables, output_docx_path)

if __name__ == "__main__":
    input_pdf_path = 'input.pdf'
    output_docx_path = 'output.docx'
    convert_pdf_to_docx(input_pdf_path, output_docx_path)
